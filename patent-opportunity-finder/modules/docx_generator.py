"""
DOCX Patent Generator Module
============================
Generate USPTO-compliant .docx provisional patent applications.

Features:
- Proper margins (1.5" left, 1" others) per USPTO guidelines
- Times New Roman 12pt, 1.5 line spacing
- Embedded images with captions
- All 9 patent sections
- Reference numeral tracking
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dataclasses import dataclass, field
from typing import List, Dict, Optional
import os
import io
from datetime import datetime
import re


@dataclass
class PatentFigure:
    """Represents a patent figure with metadata"""
    figure_number: int
    title: str
    description: str
    image_data: bytes  # Raw image bytes
    image_format: str  # png, jpg, svg
    reference_numerals: Dict[int, str] = field(default_factory=dict)


@dataclass
class InventorInfo:
    """Inventor information for cover sheet"""
    name: str
    address: str
    city: str
    state: str
    zip_code: str
    country: str = "United States"
    entity_type: str = "Micro Entity"  # Micro Entity, Small Entity, Large Entity


@dataclass
class PatentDocument:
    """Complete patent document structure"""
    title: str
    inventor: InventorInfo
    field_of_invention: str
    background: str
    summary: str
    brief_description_drawings: str
    detailed_description: str
    claims: List[str]
    abstract: str
    figures: List[PatentFigure]
    assignee: Optional[str] = None
    filing_date: datetime = field(default_factory=datetime.now)


class DocxPatentGenerator:
    """
    Generate USPTO-compliant .docx provisional patent applications.

    USPTO Formatting Standards:
    - Margins: 1.5" left, 1" top/bottom/right
    - Font: Times New Roman 12pt
    - Line spacing: 1.5 or double
    - Page size: Letter (8.5" x 11")
    """

    # USPTO formatting standards
    MARGIN_TOP = Inches(1.0)
    MARGIN_BOTTOM = Inches(1.0)
    MARGIN_LEFT = Inches(1.5)
    MARGIN_RIGHT = Inches(1.0)
    FONT_NAME = "Times New Roman"
    FONT_SIZE_BODY = Pt(12)
    FONT_SIZE_HEADING1 = Pt(14)
    FONT_SIZE_HEADING2 = Pt(13)
    LINE_SPACING = 1.5

    def __init__(self):
        self.doc = None
        self._reference_numerals = {}

    def _init_document(self):
        """Initialize document with USPTO-compliant formatting"""
        self.doc = Document()

        # Set page margins
        for section in self.doc.sections:
            section.top_margin = self.MARGIN_TOP
            section.bottom_margin = self.MARGIN_BOTTOM
            section.left_margin = self.MARGIN_LEFT
            section.right_margin = self.MARGIN_RIGHT
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        # Configure default Normal style
        style = self.doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = self.FONT_SIZE_BODY
        style.paragraph_format.line_spacing = self.LINE_SPACING
        style.paragraph_format.space_after = Pt(6)

        # Set font for East Asian text
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)

        # Configure Heading 1
        h1_style = self.doc.styles['Heading 1']
        h1_style.font.name = self.FONT_NAME
        h1_style.font.size = self.FONT_SIZE_HEADING1
        h1_style.font.bold = True
        h1_style.font.all_caps = True

        # Configure Heading 2
        h2_style = self.doc.styles['Heading 2']
        h2_style.font.name = self.FONT_NAME
        h2_style.font.size = self.FONT_SIZE_HEADING2
        h2_style.font.bold = True

    def generate(self, patent: PatentDocument, output_path: str) -> str:
        """
        Generate complete .docx from PatentDocument.

        Args:
            patent: PatentDocument with all sections
            output_path: Path to save .docx file

        Returns:
            Path to generated file
        """
        self._init_document()
        self._reference_numerals = {}

        # Build the document section by section

        # 1. Cover Sheet
        self._add_cover_sheet(patent)

        # 2. Title of Invention
        self._add_section_heading("TITLE OF INVENTION")
        self._add_paragraph(patent.title, bold=True, center=True)
        self._add_paragraph("")

        # 3. Field of Invention
        self._add_section_heading("FIELD OF INVENTION")
        self._add_content_paragraphs(patent.field_of_invention)

        # 4. Background
        self._add_section_heading("BACKGROUND")
        self._add_content_paragraphs(patent.background)

        # 5. Summary of Invention
        self._add_section_heading("SUMMARY OF INVENTION")
        self._add_content_paragraphs(patent.summary)

        # 6. Brief Description of Drawings
        self._add_section_heading("BRIEF DESCRIPTION OF DRAWINGS")
        self._add_brief_description(patent.figures)

        # 7. Detailed Description (largest section)
        self._add_section_heading("DETAILED DESCRIPTION")
        self._add_detailed_description(patent.detailed_description)

        # 8. Claims
        self._add_section_heading("CLAIMS")
        self._add_claims(patent.claims)

        # 9. Abstract
        self._add_section_heading("ABSTRACT")
        self._add_abstract(patent.abstract)

        # 10. Drawings (figures at end)
        self._add_drawings_section(patent.figures)

        # Save document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path

    def _add_cover_sheet(self, patent: PatentDocument):
        """Add USPTO cover sheet"""
        # Main title
        title = self.doc.add_paragraph()
        title_run = title.add_run("PROVISIONAL PATENT APPLICATION")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_paragraph("")

        # Application info table
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(4.5)

        fields = [
            ("Application Type:", "Provisional Patent Application"),
            ("Title:", patent.title),
            ("Filing Date:", patent.filing_date.strftime("%B %d, %Y")),
            ("Inventor:", patent.inventor.name),
            ("Address:", f"{patent.inventor.address}"),
            ("City/State/Zip:", f"{patent.inventor.city}, {patent.inventor.state} {patent.inventor.zip_code}"),
            ("Entity Status:", patent.inventor.entity_type),
            ("Assignee:", patent.assignee or "Individual Inventor"),
        ]

        for i, (label, value) in enumerate(fields):
            label_cell = table.rows[i].cells[0]
            value_cell = table.rows[i].cells[1]

            label_para = label_cell.paragraphs[0]
            label_para.add_run(label).bold = True

            value_cell.paragraphs[0].add_run(value)

        self._add_paragraph("")

        # Filing fee info
        fee_text = "Filing Fee: "
        if patent.inventor.entity_type == "Micro Entity":
            fee_text += "$160 (Micro Entity)"
        elif patent.inventor.entity_type == "Small Entity":
            fee_text += "$320 (Small Entity)"
        else:
            fee_text += "$640 (Large Entity)"

        self._add_paragraph(fee_text)

        # Page break after cover sheet
        self.doc.add_page_break()

    def _add_section_heading(self, title: str):
        """Add a section heading"""
        heading = self.doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_paragraph(self, text: str, bold: bool = False, center: bool = False, indent: bool = False):
        """Add a paragraph with optional formatting"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold

        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if indent:
            para.paragraph_format.first_line_indent = Inches(0.5)

        return para

    def _add_content_paragraphs(self, content: str):
        """Add content split into paragraphs with proper formatting"""
        # Split on double newlines for paragraphs
        paragraphs = content.split('\n\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Check for markdown headers
                if para_text.startswith('###'):
                    # Subheading level 3
                    self.doc.add_heading(para_text.replace('#', '').strip(), level=3)
                elif para_text.startswith('##'):
                    # Subheading level 2
                    self.doc.add_heading(para_text.replace('#', '').strip(), level=2)
                elif para_text.startswith('#'):
                    # Subheading level 1
                    self.doc.add_heading(para_text.replace('#', '').strip(), level=2)
                elif para_text.startswith('- ') or para_text.startswith('* '):
                    # Bullet list
                    items = para_text.split('\n')
                    for item in items:
                        item = item.lstrip('- *').strip()
                        if item:
                            self.doc.add_paragraph(item, style='List Bullet')
                else:
                    # Regular paragraph
                    self._add_paragraph(para_text, indent=True)

    def _add_brief_description(self, figures: List[PatentFigure]):
        """Add brief description of drawings section"""
        for fig in figures:
            text = f"FIG. {fig.figure_number} {fig.description}"
            self._add_paragraph(text, indent=True)

    def _add_detailed_description(self, content: str):
        """Add detailed description with proper subheadings and formatting"""
        # Split content by lines
        lines = content.split('\n')
        current_para_lines = []

        for line in lines:
            line_stripped = line.strip()

            if line_stripped.startswith('###'):
                # Flush current paragraph
                if current_para_lines:
                    self._add_paragraph(' '.join(current_para_lines), indent=True)
                    current_para_lines = []
                # Add subheading
                self.doc.add_heading(line_stripped.replace('#', '').strip(), level=3)

            elif line_stripped.startswith('##'):
                if current_para_lines:
                    self._add_paragraph(' '.join(current_para_lines), indent=True)
                    current_para_lines = []
                self.doc.add_heading(line_stripped.replace('#', '').strip(), level=2)

            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                if current_para_lines:
                    self._add_paragraph(' '.join(current_para_lines), indent=True)
                    current_para_lines = []
                item = line_stripped.lstrip('- *').strip()
                self.doc.add_paragraph(item, style='List Bullet')

            elif line_stripped:
                current_para_lines.append(line_stripped)

            elif current_para_lines:
                # Empty line - flush paragraph
                self._add_paragraph(' '.join(current_para_lines), indent=True)
                current_para_lines = []

        # Flush remaining
        if current_para_lines:
            self._add_paragraph(' '.join(current_para_lines), indent=True)

    def _add_claims(self, claims: List[str]):
        """Add claims section"""
        intro = self.doc.add_paragraph()
        intro.add_run("What is claimed is:").italic = True

        self._add_paragraph("")

        for i, claim in enumerate(claims, 1):
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.3)

            # Claim number
            num_run = para.add_run(f"{i}. ")
            num_run.bold = True

            # Claim text
            para.add_run(claim)

            self._add_paragraph("")

    def _add_abstract(self, abstract: str):
        """Add abstract section (150 words max)"""
        # Enforce word limit
        words = abstract.split()
        if len(words) > 150:
            abstract = ' '.join(words[:150]) + "..."
            warning = self.doc.add_paragraph()
            warning.add_run("[Note: Abstract truncated to 150 words per USPTO requirements]").italic = True

        self._add_paragraph(abstract, indent=True)

    def _add_drawings_section(self, figures: List[PatentFigure]):
        """Add drawings section with embedded images"""
        if not figures:
            return

        self.doc.add_page_break()
        self._add_section_heading("DRAWINGS")

        for fig in figures:
            # Figure heading
            fig_heading = self.doc.add_heading(
                f"FIG. {fig.figure_number} - {fig.title}",
                level=2
            )
            fig_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Try to embed image
            try:
                if fig.image_data and len(fig.image_data) > 0:
                    if fig.image_format.lower() == 'svg':
                        # SVG placeholder - needs conversion
                        placeholder = self.doc.add_paragraph()
                        placeholder.add_run(f"[SVG Figure - See attached file: FIG_{fig.figure_number}.svg]").italic = True
                        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Embed PNG/JPG
                        image_stream = io.BytesIO(fig.image_data)
                        self.doc.add_picture(image_stream, width=Inches(6.0))

                        # Center the image
                        last_paragraph = self.doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    placeholder = self.doc.add_paragraph()
                    placeholder.add_run(f"[Image placeholder: {fig.title}]").italic = True
                    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER

            except Exception as e:
                placeholder = self.doc.add_paragraph()
                placeholder.add_run(f"[Image: {fig.title} - Error: {str(e)}]").italic = True
                placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Caption
            caption = self.doc.add_paragraph()
            caption.add_run(fig.description).italic = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Reference numerals legend (if any)
            if fig.reference_numerals:
                self._add_paragraph("")
                legend_heading = self.doc.add_paragraph()
                legend_heading.add_run("Reference Numerals:").bold = True

                for num, desc in sorted(fig.reference_numerals.items()):
                    self.doc.add_paragraph(f"{num} - {desc}", style='List Bullet')

            self._add_paragraph("")

    def generate_from_drafter(
        self,
        provisional_patent,  # ProvisionalPatent from patent_drafter.py
        inventor: InventorInfo,
        figures: List[PatentFigure],
        output_path: str,
        assignee: str = None
    ) -> str:
        """
        Generate .docx from existing ProvisionalPatent object.

        Args:
            provisional_patent: ProvisionalPatent from patent_drafter module
            inventor: InventorInfo for cover sheet
            figures: List of PatentFigure objects
            output_path: Path to save .docx
            assignee: Optional company/organization name

        Returns:
            Path to generated file
        """
        # Convert ProvisionalPatent to PatentDocument
        patent_doc = PatentDocument(
            title=provisional_patent.title,
            inventor=inventor,
            field_of_invention=provisional_patent.field,
            background=provisional_patent.background,
            summary=provisional_patent.summary,
            brief_description_drawings=provisional_patent.brief_description_drawings,
            detailed_description=provisional_patent.detailed_description,
            claims=provisional_patent.claims,
            abstract=provisional_patent.abstract,
            figures=figures,
            assignee=assignee
        )

        return self.generate(patent_doc, output_path)


# Convenience function
def generate_patent_docx(
    patent_data: dict,
    inventor_info: dict,
    figures: List[dict],
    output_path: str
) -> str:
    """
    Convenience function to generate patent .docx from dictionaries.

    Args:
        patent_data: Dictionary with patent sections
        inventor_info: Dictionary with inventor details
        figures: List of figure dictionaries
        output_path: Output file path

    Returns:
        Path to generated file
    """
    # Create inventor
    inventor = InventorInfo(
        name=inventor_info.get('name', 'Inventor Name'),
        address=inventor_info.get('address', '123 Main St'),
        city=inventor_info.get('city', 'City'),
        state=inventor_info.get('state', 'ST'),
        zip_code=inventor_info.get('zip_code', '12345'),
        country=inventor_info.get('country', 'United States'),
        entity_type=inventor_info.get('entity_type', 'Micro Entity')
    )

    # Create figures
    patent_figures = []
    for fig in figures:
        patent_figures.append(PatentFigure(
            figure_number=fig.get('figure_number', 1),
            title=fig.get('title', 'Figure'),
            description=fig.get('description', ''),
            image_data=fig.get('image_data', b''),
            image_format=fig.get('image_format', 'png'),
            reference_numerals=fig.get('reference_numerals', {})
        ))

    # Create patent document
    patent = PatentDocument(
        title=patent_data.get('title', 'Invention Title'),
        inventor=inventor,
        field_of_invention=patent_data.get('field', ''),
        background=patent_data.get('background', ''),
        summary=patent_data.get('summary', ''),
        brief_description_drawings=patent_data.get('brief_description_drawings', ''),
        detailed_description=patent_data.get('detailed_description', ''),
        claims=patent_data.get('claims', []),
        abstract=patent_data.get('abstract', ''),
        figures=patent_figures,
        assignee=patent_data.get('assignee')
    )

    # Generate
    generator = DocxPatentGenerator()
    return generator.generate(patent, output_path)


if __name__ == "__main__":
    # Test the module
    print("Testing DOCX Patent Generator...")

    # Sample patent data
    inventor = InventorInfo(
        name="John Doe",
        address="123 Innovation Way",
        city="Tech City",
        state="CA",
        zip_code="94000",
        country="United States",
        entity_type="Micro Entity"
    )

    sample_figures = [
        PatentFigure(
            figure_number=1,
            title="System Architecture",
            description="illustrates a system 100 for processing data according to various embodiments.",
            image_data=b'',  # Empty for test
            image_format='png',
            reference_numerals={100: "System", 101: "Client", 102: "Server"}
        )
    ]

    patent = PatentDocument(
        title="System and Method for Improved Data Processing",
        inventor=inventor,
        field_of_invention="The present invention relates generally to data processing systems.",
        background="Conventional data processing systems suffer from latency issues.",
        summary="The present invention provides an improved data processing system.",
        brief_description_drawings="FIG. 1 illustrates a system architecture.",
        detailed_description="## Overview\n\nThe following describes the invention in detail...",
        claims=["A method for processing data, comprising: receiving input; processing the input; outputting results."],
        abstract="A system and method for improved data processing is disclosed.",
        figures=sample_figures
    )

    generator = DocxPatentGenerator()
    output = generator.generate(patent, "test_patent.docx")
    print(f"Generated: {output}")
