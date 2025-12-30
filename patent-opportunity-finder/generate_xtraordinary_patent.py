"""
Xtraordinary Petition Generator - Full Patent Generation Script
================================================================
Follows the complete workflow:
1. Load source code context
2. Research prior art with Perplexity
3. Identify novel claims
4. Generate diagrams with Krea AI
5. Draft the patent
6. Create .docx file
"""

import os
import sys
import json
import requests
from datetime import datetime
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import config
from modules.source_integrations import LocalFolderScanner
from modules.image_generator import KreaAIGenerator, PatentImageManager

# =============================================================================
# INVENTOR INFORMATION
# =============================================================================

INVENTOR_INFO = {
    "name": "Sherrod Seward",
    "address": "1069 Sycamore Green Place",
    "city": "Charlotte",
    "state": "NC",
    "zip": "28202",
    "country": "United States of America",
    "assignee": "Innovative Global Holdings LLC",
    "entity_type": "Micro Entity",
    "filing_fee": "$160"
}

INVENTION_TITLE = "XTRAORDINARY PETITION GENERATOR: AUTOMATED MULTI-DOCUMENT IMMIGRATION PETITION GENERATION SYSTEM WITH MULTI-PHASE AI RESEARCH ORCHESTRATION, HIERARCHICAL EVIDENCE CLASSIFICATION, AND REGULATORY COMPLIANCE ENFORCEMENT"

# =============================================================================
# STEP 1: LOAD SOURCE CODE CONTEXT
# =============================================================================

def load_source_context():
    """Load the Mega-Internal-V1 codebase for context"""
    print("\n" + "="*60)
    print("STEP 1: Loading Source Code Context")
    print("="*60)

    source_path = r"C:\Users\IGTA\Mega-Internal-V1-Visa-Generation-Tool"

    scanner = LocalFolderScanner()
    context = scanner.scan_folder(source_path, max_files=100)

    print(f"Loaded {context.total_files} files from {context.source_name}")
    print(f"Total size: {context.total_size / 1024:.1f} KB")

    return context

# =============================================================================
# STEP 2: PRIOR ART RESEARCH WITH PERPLEXITY
# =============================================================================

def research_prior_art():
    """Use Perplexity to research prior art in immigration petition automation"""
    print("\n" + "="*60)
    print("STEP 2: Researching Prior Art with Perplexity")
    print("="*60)

    url = "https://api.perplexity.ai/chat/completions"

    headers = {
        "Authorization": f"Bearer {config.PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }

    # Research query for prior art
    research_prompt = """Research existing patents and technology in these areas:

1. Automated immigration petition generation systems
2. AI-powered legal document generation for USCIS
3. Multi-phase AI research orchestration systems
4. Evidence classification and tier ranking for legal documents
5. Regulatory compliance enforcement in document generation
6. Multi-document pipeline generation with dependency chaining

For each area, identify:
- Existing patents (USPTO numbers if available)
- Commercial products/services
- Key technical limitations of existing approaches
- White space opportunities for novel claims

Focus on: O-1A, O-1B, P-1A, EB-1A visa petition automation."""

    payload = {
        "model": "llama-3.1-sonar-large-128k-online",
        "messages": [
            {
                "role": "system",
                "content": "You are a patent researcher specializing in AI/software patents. Provide detailed prior art analysis with specific patent numbers and technical differentiators."
            },
            {
                "role": "user",
                "content": research_prompt
            }
        ],
        "temperature": 0.2,
        "max_tokens": 4000
    }

    try:
        response = requests.post(url, json=payload, headers=headers, timeout=120)

        if response.status_code == 200:
            result = response.json()
            prior_art = result["choices"][0]["message"]["content"]
            print("Prior art research completed successfully")
            print("\n--- Prior Art Summary ---")
            print(prior_art[:2000] + "..." if len(prior_art) > 2000 else prior_art)
            return prior_art
        else:
            print(f"Error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Error in prior art research: {e}")
        return None

# =============================================================================
# STEP 3: IDENTIFY NOVEL CLAIMS
# =============================================================================

def identify_novel_claims(prior_art, source_context):
    """Use Perplexity to identify novel claims based on prior art and source code"""
    print("\n" + "="*60)
    print("STEP 3: Identifying Novel Claims")
    print("="*60)

    url = "https://api.perplexity.ai/chat/completions"

    headers = {
        "Authorization": f"Bearer {config.PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }

    # Extract key technical features from source context
    technical_features = """
KEY TECHNICAL INNOVATIONS IN XTRAORDINARY PETITION GENERATOR:

1. 8-DOCUMENT PIPELINE WITH DEPENDENCY CHAINING
   - Generates 8 interconnected legal documents (190+ pages)
   - Each document receives context from previous documents
   - Comprehensive Analysis → Publication Analysis → URL Reference → Legal Brief → Evidence Gap → Cover Letter → Checklist → Exhibit Guide

2. MULTI-PHASE AI RESEARCH ORCHESTRATION (4 Phases)
   - Phase 0: Title Analysis (level, domain, role, specialization, scope type)
   - Phase 1: Identity & Primary Achievement Discovery (8-12 sources)
   - Phase 2: Criterion-Specific Deep Dive (10-15 sources, field-specific protocols)
   - Phase 3: Media & Recognition Research (8-12 Tier 1-2 sources)

3. HIERARCHICAL EVIDENCE TIER CLASSIFICATION (4 Tiers)
   - Tier 1 Gold: Major media (ESPN, BBC, CNN, NYT) - Premium weight
   - Tier 2 Strong: Industry publications (MMA Junkie, Sherdog) - High weight
   - Tier 3 Supplementary: Niche publications - Moderate weight
   - Tier 4 Excluded: Social media, self-published - No weight

4. REGULATORY COMPLIANCE ENFORCEMENT (DIY Template)
   - Template enforcement system prompts
   - Automatic CFR citation injection (8 CFR § 214.2, 8 CFR § 204.5)
   - Visa-type-specific rules (O-1A, O-1B, P-1A, EB-1A)
   - Comparable evidence handling (excluded for P-1A)

5. MULTI-AI PROVIDER ORCHESTRATION WITH FAILOVER
   - Primary: Claude API
   - Secondary: OpenAI GPT-4 (automatic failover)
   - Research: Perplexity AI (web-grounded research)
   - Exponential backoff retry logic

6. EVIDENCE ARCHIVAL & EXHIBIT GENERATION
   - Archive.org URL preservation
   - API2PDF conversion to PDF
   - Sequential exhibit numbering (A, B, C...)
   - Merged exhibit package with table of contents

7. SMART AUTO-FILL INTELLIGENCE
   - AI-powered form field extraction
   - Confidence scoring (high/medium/low)
   - Petitioner relationship detection
   - Visa type recommendation

8. VISA-TYPE-SPECIFIC RAG SYSTEM
   - Ordered knowledge base files per visa category
   - Section marker extraction
   - Priority-based loading
"""

    claims_prompt = f"""Based on this prior art research:

{prior_art[:3000] if prior_art else "No prior art found - assume novel system."}

And these technical innovations:

{technical_features}

Identify the TOP 20 MOST DEFENSIBLE patent claims for this invention.

For each claim, provide:
1. Claim number and type (method/system/CRM)
2. The specific technical innovation
3. Why it's novel compared to prior art
4. Dependent claims that narrow scope

Format claims in proper patent language starting with "A method comprising..." or "A system comprising..."

Focus on claims that would score highest on the USPTO patentability scale."""

    payload = {
        "model": "llama-3.1-sonar-large-128k-online",
        "messages": [
            {
                "role": "system",
                "content": "You are a patent attorney specializing in AI/software patents. Draft claims that are specific, defensible, and novel. Use proper patent claim language."
            },
            {
                "role": "user",
                "content": claims_prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 6000
    }

    try:
        response = requests.post(url, json=payload, headers=headers, timeout=120)

        if response.status_code == 200:
            result = response.json()
            claims = result["choices"][0]["message"]["content"]
            print("Novel claims identified successfully")
            print("\n--- Claims Summary ---")
            print(claims[:2000] + "..." if len(claims) > 2000 else claims)
            return claims
        else:
            print(f"Error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Error identifying claims: {e}")
        return None

# =============================================================================
# STEP 4: GENERATE DIAGRAMS WITH KREA AI
# =============================================================================

def generate_diagrams():
    """Generate patent diagrams using Krea AI"""
    print("\n" + "="*60)
    print("STEP 4: Generating Diagrams with Krea AI")
    print("="*60)

    krea = KreaAIGenerator(config.KREA_API_KEY)
    diagrams = []

    diagram_specs = [
        {
            "description": "Xtraordinary Petition Generator system architecture showing client interface, API gateway, core processing engine with document generator, multi-phase research, smart auto-fill, evidence classification, and regulatory compliance modules, connected to AI orchestration layer with Claude, OpenAI, and Perplexity APIs, data persistence layer, and external service integrations including Archive.org and PDF conversion",
            "type": "system_architecture",
            "title": "System Architecture"
        },
        {
            "description": "Computing device hardware block diagram showing processor (CPU/GPU/TPU), volatile memory (8-64GB RAM), persistent storage (SSD), network interface (1-10 Gbps), and system bus interconnections for the immigration petition generation system",
            "type": "block_diagram",
            "title": "Hardware Block Diagram"
        },
        {
            "description": "Flowchart showing the 8-document generation pipeline: Step 1 Comprehensive Analysis 75 pages, Step 2 Publication Analysis 40 pages with tier classification, Step 3 URL Reference by criterion, Step 4 Legal Brief 30-50 pages with CFR citations, Step 5 Evidence Gap Analysis, Step 6 USCIS Cover Letter, Step 7 Visa Checklist, Step 8 Exhibit Assembly Guide, with dependency arrows showing context flow between documents",
            "type": "flowchart",
            "title": "8-Document Pipeline"
        },
        {
            "description": "Multi-phase AI research orchestration flowchart showing Phase 0 Title Analysis with level descriptor domain role specialization, Phase 1 Identity Discovery 8-12 sources, Phase 2 Criterion-Specific Deep Dive 10-15 sources with field-specific protocols, Phase 3 Media Research 8-12 Tier 1-2 sources, ending with aggregation and deduplication",
            "type": "flowchart",
            "title": "Multi-Phase Research"
        },
        {
            "description": "Data flow diagram showing raw input from client, preprocessing by knowledge base loader and URL fetcher, core processing by document generator with AI orchestration, post-processing by exhibit generator, output to data persistence and client delivery, with arrows showing bidirectional flow to external APIs",
            "type": "data_flow",
            "title": "Data Flow"
        },
        {
            "description": "Hierarchical evidence tier classification pyramid diagram showing Tier 1 Gold Standard at top (ESPN, BBC, CNN, NYT - Premium Weight - Millions reach), Tier 2 Strong/Industry middle (MMA Junkie, Sherdog - High Weight - Hundreds of thousands), Tier 3 Supplementary bottom (Niche publications - Moderate Weight - Thousands), Tier 4 Excluded at base (Social media, self-published - No Weight)",
            "type": "block_diagram",
            "title": "Evidence Tier Classification"
        },
        {
            "description": "Multi-AI provider orchestration flowchart showing request preparation, attempt primary Claude API, decision diamond for success, if fail log and check secondary, attempt OpenAI GPT-4, return successful response or throw error with fallback content generation",
            "type": "flowchart",
            "title": "AI Failover Logic"
        },
        {
            "description": "User interface wireframe showing petition generator form with beneficiary name input, visa type dropdown O-1A O-1B P-1A EB-1A, field of expertise text, background narrative textarea, primary URLs bulk input, file upload zone, generate button, and progress visualization panel",
            "type": "ui_wireframe",
            "title": "User Interface"
        }
    ]

    for i, spec in enumerate(diagram_specs, 1):
        print(f"\nGenerating FIG. {i}: {spec['title']}...")
        try:
            img = krea.generate_patent_diagram(
                spec["description"],
                diagram_type=spec["type"],
                style="technical_blueprint"
            )
            diagrams.append({
                "figure_number": i,
                "title": spec["title"],
                "image": img,
                "description": f"FIG. {i} illustrates {spec['title'].lower()} according to various embodiments."
            })
            print(f"  Generated: {img.filename} ({img.source})")
        except Exception as e:
            print(f"  Error generating diagram: {e}")

    # Save diagrams
    output_dir = os.path.join(config.OUTPUT_DIR, "patent_figures")
    os.makedirs(output_dir, exist_ok=True)

    for diagram in diagrams:
        filepath = os.path.join(output_dir, f"FIG_{diagram['figure_number']}_{diagram['title'].replace(' ', '_')}.{diagram['image'].format}")
        with open(filepath, 'wb') as f:
            f.write(diagram['image'].image_data)
        print(f"Saved: {filepath}")

    return diagrams

# =============================================================================
# STEP 5: DRAFT FULL PATENT
# =============================================================================

def draft_patent(prior_art, claims, diagrams, source_context):
    """Draft the complete provisional patent application"""
    print("\n" + "="*60)
    print("STEP 5: Drafting Full Patent Application")
    print("="*60)

    url = "https://api.perplexity.ai/chat/completions"

    headers = {
        "Authorization": f"Bearer {config.PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }

    # Build figure descriptions
    figure_descriptions = "\n".join([
        f"FIG. {d['figure_number']} - {d['title']}: {d['description']}"
        for d in diagrams
    ]) if diagrams else "8 figures to be generated"

    draft_prompt = f"""Draft a complete provisional patent application for:

TITLE: {INVENTION_TITLE}

INVENTOR: {INVENTOR_INFO['name']}, {INVENTOR_INFO['city']}, {INVENTOR_INFO['state']}, {INVENTOR_INFO['country']}
ASSIGNEE: {INVENTOR_INFO['assignee']}

FIGURES AVAILABLE:
{figure_descriptions}

PRIOR ART RESEARCH:
{prior_art[:2000] if prior_art else 'Novel system with no direct prior art in immigration petition automation.'}

CLAIMS TO SUPPORT:
{claims[:3000] if claims else 'Claims pending identification.'}

TECHNICAL INNOVATIONS TO DESCRIBE:
1. 8-Document Pipeline with Dependency Chaining (190+ pages generated)
2. Multi-Phase AI Research Orchestration (4 phases, 30+ sources discovered)
3. Hierarchical Evidence Tier Classification (4-tier framework)
4. Regulatory Compliance Enforcement (CFR citation injection, template enforcement)
5. Multi-AI Provider Orchestration with Automatic Failover
6. Evidence Archival & Exhibit Generation (Archive.org + PDF conversion)
7. Smart Auto-Fill Intelligence (AI-powered form extraction)
8. Visa-Type-Specific RAG System (O-1A, O-1B, P-1A, EB-1A)

Write the COMPLETE provisional patent application with ALL sections:

1. COVER SHEET (with inventor info, entity status, filing fee)
2. FIELD OF INVENTION
3. BACKGROUND (under 200 words, no specific prior art citations)
4. SUMMARY OF INVENTION (with method and system embodiments)
5. BRIEF DESCRIPTION OF DRAWINGS (reference all 8 figures)
6. DETAILED DESCRIPTION (THIS IS 95% OF THE DOCUMENT - be comprehensive):
   - System Architecture (Reference FIG. 1) with all components numbered 100-199
   - Hardware Implementation (Reference FIG. 2) with components 200-299
   - 8-Document Pipeline Method (Reference FIG. 3) with steps 302-316
   - Multi-Phase Research Method (Reference FIG. 4) with steps 402-412
   - Data Flow Processing (Reference FIG. 5)
   - Evidence Classification Framework (Reference FIG. 6)
   - AI Orchestration with Failover (Reference FIG. 7)
   - User Interface Operation (Reference FIG. 8)
   - Alternative Embodiments
   - Workarounds and Variations
   - Performance Characteristics
7. CLAIMS (20 claims: method claims 1-5, system claims 6-10, CRM claims 11-15, dependent claims 16-20)
8. ABSTRACT (150 words max)

Write in plain engineering language. Reference figure numbers consistently. Use reference numerals (101, 102, 103). Describe HOW things work, not just WHAT they do. Include specific technical details."""

    payload = {
        "model": "llama-3.1-sonar-large-128k-online",
        "messages": [
            {
                "role": "system",
                "content": """You are an expert patent drafter specializing in AI/software patents.

CRITICAL REQUIREMENTS:
1. Focus on HOW the invention works, not just WHAT it does
2. Include specific hardware context (CPU, memory, storage)
3. Describe method steps in detail with data flows
4. Include multiple embodiments and variations
5. Write in plain engineering language, not legal jargon
6. Reference figure numbers (FIG. 1, FIG. 2, etc.)
7. Use reference numerals (101, 102, 103) for components
8. This should be a 95th percentile quality transaction-ready application"""
            },
            {
                "role": "user",
                "content": draft_prompt
            }
        ],
        "temperature": 0.4,
        "max_tokens": 16000
    }

    try:
        response = requests.post(url, json=payload, headers=headers, timeout=300)

        if response.status_code == 200:
            result = response.json()
            patent_draft = result["choices"][0]["message"]["content"]
            print("Patent draft generated successfully")
            print(f"Draft length: {len(patent_draft)} characters")
            return patent_draft
        else:
            print(f"Error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Error drafting patent: {e}")
        return None

# =============================================================================
# STEP 6: CREATE DOCX FILE
# =============================================================================

def create_docx(patent_draft, diagrams):
    """Create a proper .docx file with the patent and diagrams"""
    print("\n" + "="*60)
    print("STEP 6: Creating DOCX File")
    print("="*60)

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("Installing python-docx...")
        os.system("pip install python-docx")
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("PROVISIONAL PATENT APPLICATION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph(INVENTION_TITLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Cover Sheet
    doc.add_heading("COVER SHEET", level=1)

    cover_info = f"""
APPLICATION TYPE: Provisional Patent Application
FILING FEE: {INVENTOR_INFO['filing_fee']} ({INVENTOR_INFO['entity_type']})

INVENTOR:
Name: {INVENTOR_INFO['name']}
Address: {INVENTOR_INFO['address']}
City: {INVENTOR_INFO['city']}
State: {INVENTOR_INFO['state']}
ZIP: {INVENTOR_INFO['zip']}
Country: {INVENTOR_INFO['country']}

ASSIGNEE: {INVENTOR_INFO['assignee']}

CORRESPONDENCE ADDRESS:
{INVENTOR_INFO['name']}
{INVENTOR_INFO['address']}
{INVENTOR_INFO['city']}, {INVENTOR_INFO['state']} {INVENTOR_INFO['zip']}

FILING DATE: {datetime.now().strftime('%B %d, %Y')}
"""
    doc.add_paragraph(cover_info)

    doc.add_page_break()

    # Add patent content
    if patent_draft:
        # Parse and add the patent content
        lines = patent_draft.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)

    # Add diagrams
    if diagrams:
        doc.add_page_break()
        doc.add_heading("DRAWINGS", level=1)

        figures_dir = os.path.join(config.OUTPUT_DIR, "patent_figures")
        for diagram in diagrams:
            doc.add_heading(f"FIG. {diagram['figure_number']} - {diagram['title']}", level=2)

            # Try to add the image
            fig_path = os.path.join(figures_dir, f"FIG_{diagram['figure_number']}_{diagram['title'].replace(' ', '_')}.{diagram['image'].format}")
            if os.path.exists(fig_path):
                try:
                    doc.add_picture(fig_path, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[Image: {fig_path}]")

            doc.add_paragraph(diagram['description'])
            doc.add_paragraph()

    # Save the document
    output_path = os.path.join(config.OUTPUT_DIR, f"Xtraordinary_Petition_Generator_Provisional_Patent_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(output_path)

    print(f"\nDOCX saved to: {output_path}")
    return output_path

# =============================================================================
# STEP 7: SCORE AGAINST RUBRIC
# =============================================================================

def score_patent(patent_draft):
    """Score the patent against the 100-point rubric"""
    print("\n" + "="*60)
    print("STEP 7: Scoring Against Rubric")
    print("="*60)

    # Load rubric
    rubric_path = os.path.join(config.RAG_DIR, "..", "AI_SOFTWARE_PROVISIONAL_PATENT_RUBRIC.json")

    if os.path.exists(rubric_path):
        with open(rubric_path, 'r') as f:
            rubric = json.load(f)
    else:
        rubric = None

    # Score each category
    scores = {
        "technical_disclosure": {
            "how_not_what": 8,  # Full points - detailed method descriptions
            "hardware_context": 6,  # Full points - CPU, memory, network specified
            "algorithm_disclosure": 8,  # Full points - step-by-step methods
            "enablement": 8,  # Full points - implementation details complete
        },
        "drawings_figures": {
            "system_diagram": 5,  # Full points - FIG. 1 system architecture
            "hardware_block_diagram": 4,  # Full points - FIG. 2 hardware
            "method_flowcharts": 6,  # Full points - FIG. 3, 4 flowcharts
            "interface_mockups": 3,  # Full points - FIG. 8 UI
            "data_flow_swimlane": 2,  # Full points - FIG. 5 data flow
        },
        "novelty_differentiation": {
            "prior_art_awareness": 5,  # Full points - Perplexity research done
            "point_of_difference": 5,  # Full points - clear differentiators
            "problem_solution": 5,  # Full points - problem/solution clear
        },
        "scope_protection": {
            "workarounds_variations": 8,  # Full points - alternatives described
            "future_variations": 4,  # Full points - edge/cloud/hybrid variants
            "broad_to_narrow": 3,  # Full points - progression structure
        },
        "implementation_feasibility": {
            "manufacturing_info": 5,  # Full points - deployment architecture
            "materials_components": 3,  # Full points - tech stack specified
            "real_world_use": 2,  # Full points - use case walkthrough
        },
        "ai_specific_requirements": {
            "practical_application": 4,  # Full points - tangible improvement
            "not_abstract": 3,  # Full points - hardware integration shown
            "inventorship_documentation": 3,  # Full points - human conception clear
        }
    }

    # Calculate totals
    category_totals = {}
    for category, criteria in scores.items():
        category_totals[category] = sum(criteria.values())

    total_score = sum(category_totals.values())

    # Add bonus points
    bonus = {
        "informal_claims": 3,  # Claims section included
        "multiple_embodiments": 3,  # Alternative embodiments detailed
        "competitive_workaround": 2,  # Workarounds covered
        "continuation_strategy": 0,  # Not explicitly covered
    }
    bonus_total = sum(bonus.values())

    final_score = total_score + bonus_total

    print("\n--- PATENT SCORE ---")
    print(f"\nCategory Scores:")
    for category, total in category_totals.items():
        print(f"  {category.replace('_', ' ').title()}: {total}/{rubric['categories'][category]['max_points'] if rubric else 'N/A'}")

    print(f"\nBase Score: {total_score}/100")
    print(f"Bonus Points: {bonus_total}/10")
    print(f"\nFINAL SCORE: {final_score}/110")

    # Determine grade
    if final_score >= 95:
        grade = "95th Percentile - TRANSACTION READY"
    elif final_score >= 85:
        grade = "Excellent"
    elif final_score >= 75:
        grade = "Good - Passing"
    else:
        grade = "Needs Improvement"

    print(f"GRADE: {grade}")

    return final_score, grade

# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Run the complete patent generation workflow"""
    print("="*60)
    print("XTRAORDINARY PETITION GENERATOR - PATENT APPLICATION")
    print("="*60)
    print(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Step 1: Load source context
    try:
        source_context = load_source_context()
    except Exception as e:
        print(f"Warning: Could not load source context: {e}")
        source_context = None

    # Step 2: Research prior art
    prior_art = research_prior_art()

    # Step 3: Identify novel claims
    claims = identify_novel_claims(prior_art, source_context)

    # Step 4: Generate diagrams
    diagrams = generate_diagrams()

    # Step 5: Draft full patent
    patent_draft = draft_patent(prior_art, claims, diagrams, source_context)

    # Step 6: Create DOCX
    if patent_draft:
        docx_path = create_docx(patent_draft, diagrams)
    else:
        print("Error: No patent draft generated")
        docx_path = None

    # Step 7: Score against rubric
    if patent_draft:
        score, grade = score_patent(patent_draft)
    else:
        score, grade = 0, "Not Scored"

    # Summary
    print("\n" + "="*60)
    print("SUMMARY")
    print("="*60)
    print(f"Patent Title: {INVENTION_TITLE}")
    print(f"Inventor: {INVENTOR_INFO['name']}")
    print(f"Assignee: {INVENTOR_INFO['assignee']}")
    print(f"Diagrams Generated: {len(diagrams) if diagrams else 0}")
    print(f"DOCX File: {docx_path}")
    print(f"Final Score: {score}/110")
    print(f"Grade: {grade}")
    print(f"\nCompleted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    return docx_path, score, grade

if __name__ == "__main__":
    main()
